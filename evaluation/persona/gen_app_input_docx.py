"""앱 입력 데이터 docx 생성 — 실제 앱 입력 흐름(Figma 디자인 + BE DTO) 형식으로.

각 data/<이름>/subject_context.json → 앱입력데이터_<이름>.docx.
흐름: 온보딩(프로필) → 동의 → 가족 용어집 → Voice Persona 8섹션 설문 → 음성 샘플.
"""

import json
from pathlib import Path

from docx import Document

REPO = Path(__file__).resolve().parents[2]
DATA = REPO / "data"
NAMES = ["신금자", "최영자", "이순덕", "정말순", "서정숙", "김분남"]

_CONSENTS = [
    ("[필수] 개인정보 수집·이용", "동의"),
    ("[필수] 음성 파일 처리", "동의"),
    ("[선택] AI 학습 활용", "동의(테스트)"),
    ("[선택] 국외 이전", "동의(테스트)"),
    ("[선택] 사후 AI 사용 (Memories·Persona 활성화)", "동의(테스트)"),
]


def _table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for r in rows:
        cells = t.add_row().cells
        for i, v in enumerate(r):
            cells[i].text = str(v if v is not None else "")


def _life_status(status: str) -> str:
    s = (status or "").strip()
    if "사망" in s or "고인" in s or s.upper() == "DECEASED":
        return "곁을 떠나셨어요 (DECEASED)"
    if "생존" in s or s.upper() == "LIVING":
        return "지금 함께 계세요 (LIVING)"
    return f"{s or '미상'} (UNKNOWN)"


def generate(name: str, payload: dict, dest: Path) -> None:
    sc = payload.get("subjectContext", {})
    ic = payload.get("intakeContext", {})
    subject = sc.get("subject") or {}
    bp = ic.get("basicProfile") or {}
    fam_members = {m.get("name"): m for m in (sc.get("familyMembers") or [])}

    doc = Document()
    doc.add_heading(f"TalkTo 앱 입력 데이터 — {name}", 0)
    doc.add_paragraph("용도: 앱 온보딩·Voice Persona 설문 테스트 입력값. "
                      "실제 앱 입력 흐름(Figma 디자인 + BE 계약) 기준으로 작성.")

    doc.add_heading("1. 온보딩 — 대상자 프로필", 1)
    _table(doc, ["항목", "입력값"], [
        ["대상자 이름", subject.get("name", "")],
        ["가족 내 호칭(relationship)", subject.get("addressTerm", "")],
        ["생존 상태(lifeStatus)", _life_status(bp.get("status", ""))],
        ["지역·방언 힌트(dialectHint)", bp.get("dialect") or bp.get("region") or "(선택)"],
    ])

    doc.add_heading("2. 동의 항목", 1)
    _table(doc, ["동의 항목", "설정"], [[c, v] for c, v in _CONSENTS])

    doc.add_heading("3. 가족 용어집 (전사 보정·화자 식별용)", 1)
    glossary = []
    for m in (sc.get("familyMembers") or []):
        terms = m.get("addressTerms") or []
        glossary.append(["PERSON", m.get("name", ""), ", ".join(terms)])
    for term in (sc.get("glossaryTerms") or []):
        glossary.append(["PLACE/OTHER", term, ""])
    _table(doc, ["유형(termType)", "용어", "부른 호칭/발음"], glossary or [["-", "-", "-"]])

    doc.add_heading("4. Voice Persona 설문 (8섹션)", 1)

    doc.add_heading("4-1. 기본 정보", 2)
    doc.add_paragraph(f"한 줄 소개: {bp.get('oneLine','') or '(없음)'}")

    doc.add_heading("4-2. 가족관계와 호칭", 2)
    fam_rows = []
    for m in (ic.get("familyMap") or []):
        nm = m.get("name", "")
        addr = ", ".join((fam_members.get(nm) or {}).get("addressTerms", []))
        fam_rows.append([nm, m.get("relation", ""), addr, m.get("tone", "")])
    _table(doc, ["이름", "관계", "대상자가 부른 호칭", "챙김 톤"], fam_rows or [["-"]*4])

    doc.add_heading("4-3. 말투 스타일", 2)
    doc.add_paragraph(f"말투 성향: {ic.get('speechStyle','') or '(없음)'}")
    examples = payload.get("_speechExamples") or []
    doc.add_paragraph("말투 예시(실제 표현):")
    for ex in (examples[:12] or ["(fixture에 없음 — 녹음에서 자동 추출)"]):
        doc.add_paragraph(str(ex), style="List Bullet")

    doc.add_heading("4-4. 성격·가치", 2)
    doc.add_paragraph(ic.get("personality", "") or "(없음)")

    doc.add_heading("4-5. 대표 기억", 2)
    _table(doc, ["제목", "내용", "태그"],
           [[c.get("title", ""), c.get("content", ""), ", ".join(c.get("tags", []) or [])]
            for c in (ic.get("memoryCards") or [])] or [["-", "-", "-"]])

    doc.add_heading("4-6. 상황별 반응", 2)
    _table(doc, ["상황", "할 법한 답", "피할 것"],
           [[s.get("situation", ""), s.get("response", ""), s.get("avoid", "")]
            for s in (ic.get("situationalReactions") or [])] or [["-", "-", "-"]])

    doc.add_heading("4-7. 금기 주제 접근 범위", 2)
    for t in (ic.get("tabooTopics") or []) or ["(없음)"]:
        doc.add_paragraph(f"{t} — 제한/금지", style="List Bullet")

    doc.add_heading("5. 음성 샘플 (고인 목소리 잘 들리는 구간)", 1)
    vsr = (ic.get("sttHints") or {}).get("voiceSampleRef")
    if vsr:
        doc.add_paragraph(f"documentId: {vsr.get('documentId','')} · "
                          f"startMs: {vsr.get('startMs','')} · endMs: {vsr.get('endMs','')}")
    else:
        doc.add_paragraph("업로드한 통화 중 고인 목소리가 선명한 구간을 선택 "
                          "(recordingId + startMs/endMs).")

    doc.save(str(dest))


def main() -> int:
    for name in NAMES:
        fx = DATA / name / "subject_context.json"
        if not fx.exists():
            print(f"건너뜀(없음): {fx}")
            continue
        payload = json.loads(fx.read_text(encoding="utf-8"))
        dest = DATA / name / f"앱입력데이터_{name}.docx"
        generate(name, payload, dest)
        # 옛 축약본 Intake docx 제거(신금자 실제 설문지는 유지)
        old = DATA / name / f"Intake_{name}.docx"
        if old.exists() and name != "신금자":
            old.unlink()
        print(f"{name}: 앱입력데이터_{name}.docx 생성")
    print(f"\n완료 → {DATA}/<이름>/")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
