"""페르소나 데이터 정리 — 인물별 폴더에 원본 통화 + Intake docx + fixture(JSON) 모으기.

data/personas/<이름>/
  ├─ 통화녹음/            원본 오디오
  ├─ Intake_<이름>.docx   설문지 형식(신금자는 실제 작성본, 외부는 fixture로 생성)
  └─ subject_context.json fixture(앱/AI가 소비하는 구조화 형태)
"""

import json
import shutil
from pathlib import Path

from docx import Document

REPO = Path(__file__).resolve().parents[2]
OUT = REPO / "data"

# (이름, fixture 경로, 오디오 소스, 실제 docx(있으면))
PERSONS = [
    ("신금자", "data/fixtures/subject_context_singeumja.json", ("dir", "data/voice_raw"),
     "../TalkTo_Voice_Persona_Intake_신금자.docx"),
    ("최영자", "evaluation/persona/fixtures/subject_context_choiyoungja_gangwon.json",
     ("pair", "speakergw712_speakergw713"), None),
    ("이순덕", "evaluation/persona/fixtures/subject_context_isunduk_gangwon.json",
     ("pair", "speakergw3139_speakergw3140"), None),
    ("정말순", "evaluation/persona/fixtures/subject_context_jeongmalsun_gangwon.json",
     ("pair", "speakergw2440_speakergw2441"), None),
    ("서정숙", "evaluation/persona/fixtures/subject_context_seojeongsuk_gyeongsang.json",
     ("pair", "speakergs4913_speakergs4914"), None),
    ("김분남", "evaluation/persona/fixtures/subject_context_kimbunnam_gyeongsang.json",
     ("pair", "speakergs4921_speakergs4922"), None),
]


def _kv_table(doc, rows, headers):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val or "")


def generate_intake_docx(name: str, payload: dict, dest: Path) -> None:
    sc = payload.get("subjectContext", {})
    ic = payload.get("intakeContext", {})
    subject = sc.get("subject") or {}
    bp = ic.get("basicProfile") or {}
    doc = Document()
    doc.add_heading("TalkTo Voice Persona — Intake 설문지", 0)
    doc.add_paragraph(f"대상자: {subject.get('name','')} · 생성: fixture 기반")

    doc.add_heading("1. 대상자 기본 정보", 1)
    _kv_table(doc, [
        ["성함", subject.get("name", "")],
        ["가족 내 호칭", subject.get("addressTerm", "")],
        ["한 줄 소개", bp.get("oneLine", "")],
        ["상태", bp.get("status", "")],
    ], ["항목", "작성란"])

    doc.add_heading("2. 가족관계 지도", 1)
    fam = ic.get("familyMap") or []
    _kv_table(doc, [[m.get("name", ""), m.get("relation", ""), m.get("tone", "")] for m in fam],
              ["이름", "관계", "이 사람에게 하는 말투·챙김"])

    doc.add_heading("3. 말투·언어 스타일", 1)
    doc.add_paragraph(ic.get("speechStyle", "") or "(없음)")

    doc.add_heading("4. 성격·가치관·삶의 원칙", 1)
    doc.add_paragraph(ic.get("personality", "") or "(없음)")

    doc.add_heading("5. 대표 기억 카드", 1)
    cards = ic.get("memoryCards") or []
    _kv_table(doc, [[c.get("title", ""), c.get("content", "")] for c in cards],
              ["제목", "내용"])

    doc.add_heading("6. 상황별 Persona 반응", 1)
    sits = ic.get("situationalReactions") or []
    _kv_table(doc, [[s.get("situation", ""), s.get("response", ""), s.get("avoid", "")] for s in sits],
              ["상황", "할 법한 답", "피할 것"])

    doc.add_heading("7. 금기 주제", 1)
    for t in (ic.get("tabooTopics") or []) or ["(없음)"]:
        doc.add_paragraph(str(t), style="List Bullet")

    doc.save(str(dest))


def main() -> int:
    from dotenv import load_dotenv
    load_dotenv(REPO / ".env")
    from evaluation.e2e_external import find_clips

    for name, fx_rel, (kind, src), real_docx in PERSONS:
        person_dir = OUT / name
        audio_dir = person_dir / "통화녹음"
        audio_dir.mkdir(parents=True, exist_ok=True)

        # 오디오 복사
        clips = sorted((REPO / src).glob("*.wav")) if kind == "dir" else sorted(find_clips(src))
        for clip in clips:
            shutil.copy2(clip, audio_dir / clip.name)

        # fixture 복사
        payload = json.loads((REPO / fx_rel).read_text(encoding="utf-8"))
        (person_dir / "subject_context.json").write_text(
            json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")

        # Intake docx
        docx_dest = person_dir / f"Intake_{name}.docx"
        if real_docx and (REPO / real_docx).exists():
            shutil.copy2(REPO / real_docx, docx_dest)
            src_note = "실제 작성본"
        else:
            generate_intake_docx(name, payload, docx_dest)
            src_note = "fixture로 생성"

        print(f"{name}: 통화 {len(clips)}건 + Intake docx({src_note}) + fixture → {person_dir}")

    print(f"\n완료 → {OUT}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
